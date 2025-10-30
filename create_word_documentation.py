#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pentru Generarea Documentației Word
Spending Tracker - Documentație Tehnică Completă
Versiune: 1.0.0
Data: Octombrie 2025

Instrucțiuni:
1. Rulează scriptul: python3 create_word_documentation.py
2. Se va genera: SPENDING_TRACKER_DOCUMENTATION.docx
3. Documentul conține placeholder-uri pentru imagini - le marcate cu [INSERT IMAGE: ...]
4. Înlocuiești placeholder-urile cu imaginile reale din aplicație
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_page_break(doc):
    """Adaugă page break"""
    doc.add_page_break()

def set_cell_background(cell, fill):
    """Setează culoarea de background a unei celule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_title_page(doc):
    """Creează pagina de titlu"""
    # Logo/Titlu
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('💰 SPENDING TRACKER')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(81, 43, 212)  # #512BD4
    
    # Subtitlu 1
    subtitle1 = doc.add_paragraph()
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle1.add_run('Sistem Integrat de Gestionare și Analiză a Cheltuielilor')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(78, 205, 196)  # #4ECDC4
    
    # Subtitlu 2
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('Documentație Tehnică Completă')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(78, 205, 196)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tabel informații
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    set_cell_background(header_cells[0], '512BD4')
    set_cell_background(header_cells[1], '512BD4')
    
    # Informații
    info_data = [
        ('Versiune:', '1.0.0'),
        ('Platform:', '.NET MAUI'),
        ('Bază de Date:', 'SQLite'),
        ('Limbă Documentație:', 'Română'),
        ('Data:', datetime.now().strftime('%d/%m/%Y')),
        ('Pagini:', '20'),
        ('Status:', 'Producție')
    ]
    
    for idx, (label, value) in enumerate(info_data):
        row = table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
        
        # Format label cell
        label_cell = row.cells[0]
        set_cell_background(label_cell, 'F0F0F0')
        label_para = label_cell.paragraphs[0]
        label_para.runs[0].font.bold = True
        label_para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('© 2025 Spending Tracker. Toate drepturile rezervate.')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run('Documentație de referință pentru dezvoltatori și utilizatori')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

def add_heading(doc, text, level=1):
    """Adaugă heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style = 'Heading 1'
        for run in heading.runs:
            run.font.color.rgb = RGBColor(81, 43, 212)
            run.font.size = Pt(16)
    elif level == 2:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(78, 205, 196)
            run.font.size = Pt(13)
    else:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(69, 183, 209)
            run.font.size = Pt(11)
    return heading

def add_normal_text(doc, text):
    """Adaugă paragraf normal"""
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def add_image_placeholder(doc, image_name, description):
    """Adaugă placeholder pentru imagine"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(f'[INSERT IMAGE: {image_name}]')
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 107, 107)  # Roșu pentru a fi vizibil
    run.font.size = Pt(10)
    
    p2 = doc.add_paragraph(description)
    p2.paragraph_format.left_indent = Inches(0.5)
    for run in p2.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

def create_toc(doc):
    """Creează table of contents"""
    doc.add_heading('CUPRINS', level=1)
    
    toc_items = [
        '1. INTRODUCERE',
        '2. DESCRIERE TEHNICĂ',
        '3. MODELUL DE DATE',
        '4. SERVICII DE BAZĂ',
        '5. INTERFACE-URI UTILIZATOR',
        '6. FUNCȚIONALITĂȚI PRINCIPALE',
        '7. FLUXURI DE DATE ȘI PROCESE',
        '8. CARACTERISTICI DE SECURITATE',
        '9. GHID UTILIZATOR',
        '10. CONCLUZII ȘI PERSPECTIVĂ VIITOARE'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10)

def create_introduction(doc):
    """Secțiunea 1: Introducere"""
    add_heading(doc, '1. INTRODUCERE', level=1)
    
    add_heading(doc, '1.1. Prezentare Generală', level=2)
    add_normal_text(doc, 
        'Spending Tracker este o aplicație mobilă multiplatformă dezvoltată în .NET MAUI care oferă '
        'o soluție completă pentru gestionarea cheltuielilor zilnice. Aplicația permite utilizatorilor '
        'să înregistreze, categorizeze, analizeze și optimizeze cheltuielile lor financiare cu o interfață '
        'intuitivă și ușor de utilizat.')
    
    add_heading(doc, 'Caracteristici principale:', level=3)
    features = [
        '✓ Înregistrare rapidă a cheltuielilor cu detalii complete',
        '✓ Categorisare automată și personalizabilă (7 categorii implicite)',
        '✓ Gestionare bugete lunare per categorie cu alerte vizuale',
        '✓ Analiză statistică avansată cu progrese vizuale',
        '✓ Rapoarte detaliate pe diferite perioade (lună, 3 luni, 6 luni, an)',
        '✓ Convertor de monede în timp real (RON, EUR, USD, GBP)',
        '✓ Interfață intuitivă și ușor de utilizat',
        '✓ Suport multiplatformă (Android, iOS, Windows, macOS)',
        '✓ Bază de date locală sigură (SQLite)',
        '✓ Validare riguroasă a datelor și gestionare erori'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '1.2. Obiectivele Proiectului', level=2)
    add_normal_text(doc,
        'Obiectivul principal al Spending Tracker este de a oferi utilizatorilor o metodă simplă, '
        'eficientă și plăcută de urmărire a cheltuielilor lor zilnice. Cu ajutorul acestei aplicații, '
        'utilizatorii pot obține insight-uri valoroase despre obiceiurile lor de cheltuire și pot planifica '
        'mai bine bugetele viitoare.')
    
    add_heading(doc, '1.3. Motivație și Context', level=2)
    add_normal_text(doc,
        'În era digitală actuală, gestionarea finanțelor personale este o provocare pentru mulți utilizatori. '
        'Plăți zilnice, abonamente și cheltuieli neașteptate se acumulează rapid. Spending Tracker vine să rezolve '
        'această problemă prin oferirea unui instrument centralizat, sigur și ușor de utilizat pentru controlul cheltuielilor.')

def create_technical_description(doc):
    """Secțiunea 2: Descriere Tehnică"""
    add_page_break(doc)
    add_heading(doc, '2. DESCRIERE TEHNICĂ', level=1)
    
    add_heading(doc, '2.1. Arhitectura Sistemului', level=2)
    add_normal_text(doc,
        'Spending Tracker utilizează o arhitectură modulară cu straturi bine definite care asigură '
        'separarea responsabilităților și ușurința în mentenanță.')
    
    # Tabel arhitectură
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Stratul'
    header_cells[1].text = 'Componente'
    header_cells[2].text = 'Responsabilități'
    
    for cell in header_cells:
        set_cell_background(cell, '512BD4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    arch_data = [
        ('UI/Prezentare', 'Pages (9 pagini XAML)', 'Afișare interfață utilizator'),
        ('Logică Afaceri', 'ViewModels, Event Handlers', 'Procesare date și logică'),
        ('Servicii', 'DatabaseService, CurrencyService', 'Operații de bază'),
        ('Acces Date', 'SQLite, SQLiteConnection', 'Persistență date')
    ]
    
    for idx, (layer, components, responsibility) in enumerate(arch_data):
        row = table.rows[idx + 1]
        row.cells[0].text = layer
        row.cells[1].text = components
        row.cells[2].text = responsibility
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'F9F9F9')
    
    add_heading(doc, '2.2. Tehnologii Utilizate', level=2)
    
    tech_data = [
        ('Framework', '.NET MAUI 9.0', 'Multi-platform App UI'),
        ('Bază Date', 'SQLite', 'Bază de date locală, portabilă'),
        ('ORM', 'sqlite-net-pcl 1.9.172', 'Mapare obiect-relație'),
        ('JSON', 'Newtonsoft.Json 13.0.3', 'Serializare date externe'),
        ('Logging', 'Microsoft.Extensions.Logging', 'Debug și logging'),
        ('UI', 'XAML', 'XML Application Markup Language')
    ]
    
    tech_table = doc.add_table(rows=len(tech_data) + 1, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_header = tech_table.rows[0].cells
    tech_header[0].text = 'Categorie'
    tech_header[1].text = 'Tehnologie'
    tech_header[2].text = 'Descriere'
    
    for cell in tech_header:
        set_cell_background(cell, '45B7D1')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    for idx, (category, tech, description) in enumerate(tech_data):
        row = tech_table.rows[idx + 1]
        row.cells[0].text = category
        row.cells[1].text = tech
        row.cells[2].text = description
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'F9F9F9')
    
    add_heading(doc, '2.3. Structura Proiectului', level=2)
    add_normal_text(doc, 'Proiectul este organizat în următoarea structură de foldere și fișiere:')
    
    structure = """
Models/ (3 clase model)
├── Expense.cs - Cheltuiala
├── Category.cs - Categoria
└── Budget.cs - Bugetul

Pages/ (9 pagini cu XAML + Code-behind)
├── MainPage.xaml - Pagina principală
├── ExpenseListPage.xaml - Lista cheltuieli
├── BudgetPage.xaml - Gestionare bugete
├── CategoriesPage.xaml - Gestionare categorii
├── StatisticsPage.xaml - Statistici
├── ReportsPage.xaml - Rapoarte
├── CurrencyConverterPage.xaml - Convertor valute
├── SettingsPage.xaml - Setări
└── AboutPage.xaml - Despre

Services/ (2 servicii de bază)
├── DatabaseService.cs - Operații DB
└── CurrencyService.cs - Conversii valute

Resources/ - Imagini, fonturi, stiluri
AppShell.xaml - Navigare aplicație
MauiProgram.cs - Configurare inițială
    """
    
    p = doc.add_paragraph(structure)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def create_data_model(doc):
    """Secțiunea 3: Modelul de Date"""
    add_page_break(doc)
    add_heading(doc, '3. MODELUL DE DATE', level=1)
    
    add_heading(doc, '3.1. Entitatea Cheltuială (Expense)', level=2)
    add_normal_text(doc, 'Stochează informațiile despre fiecare cheltuiala introdusă de utilizator.')
    
    expense_table = doc.add_table(rows=7, cols=3)
    expense_table.style = 'Light Grid Accent 1'
    
    expense_header = expense_table.rows[0].cells
    expense_header[0].text = 'Atribut'
    expense_header[1].text = 'Tip'
    expense_header[2].text = 'Descriere'
    
    for cell in expense_header:
        set_cell_background(cell, 'FF6B6B')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    expense_data = [
        ('Id', 'int', 'Cheie primară, auto-increment'),
        ('Description', 'string', 'Descrierea cheltuielii'),
        ('Amount', 'double', 'Suma în valuta respectivă'),
        ('Category', 'string', 'Categoria cheltuielii'),
        ('Date', 'DateTime', 'Data și ora cheltuielii'),
        ('Currency', 'string', 'Codul valutei (RON, EUR, USD, GBP)')
    ]
    
    for idx, (attr, typ, desc) in enumerate(expense_data):
        row = expense_table.rows[idx + 1]
        row.cells[0].text = attr
        row.cells[1].text = typ
        row.cells[2].text = desc
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'FFF5F5')
    
    add_heading(doc, '3.2. Entitatea Categorie (Category)', level=2)
    add_normal_text(doc, 'Definește categoriile de cheltuieli disponibile în sistem.')
    
    cat_table = doc.add_table(rows=4, cols=3)
    cat_table.style = 'Light Grid Accent 1'
    
    cat_header = cat_table.rows[0].cells
    cat_header[0].text = 'Atribut'
    cat_header[1].text = 'Tip'
    cat_header[2].text = 'Descriere'
    
    for cell in cat_header:
        set_cell_background(cell, '4ECDC4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    cat_data = [
        ('Id', 'int', 'Cheie primară, auto-increment'),
        ('Name', 'string', 'Nume categoria'),
        ('Color', 'string', 'Cod culoare HEX (ex: #FF6B6B)')
    ]
    
    for idx, (attr, typ, desc) in enumerate(cat_data):
        row = cat_table.rows[idx + 1]
        row.cells[0].text = attr
        row.cells[1].text = typ
        row.cells[2].text = desc
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'F0FFFE')
    
    add_heading(doc, '3.3. Categorii Implicite', level=3)
    
    default_cat_table = doc.add_table(rows=8, cols=2)
    default_cat_table.style = 'Light Grid Accent 1'
    
    cat_header = default_cat_table.rows[0].cells
    cat_header[0].text = 'Categorie'
    cat_header[1].text = 'Culoare'
    
    for cell in cat_header:
        set_cell_background(cell, '512BD4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    categories = [
        ('Alimente', '#FF6B6B'),
        ('Transport', '#4ECDC4'),
        ('Divertisment', '#45B7D1'),
        ('Utilități', '#FFA07A'),
        ('Sănătate', '#98D8C8'),
        ('Educație', '#F7DC6F'),
        ('Altele', '#BB8FCE')
    ]
    
    for idx, (cat, color) in enumerate(categories):
        row = default_cat_table.rows[idx + 1]
        row.cells[0].text = cat
        row.cells[1].text = color
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'F9F9F9')
    
    add_heading(doc, '3.4. Entitatea Buget (Budget)', level=2)
    
    budget_table = doc.add_table(rows=6, cols=3)
    budget_table.style = 'Light Grid Accent 1'
    
    budget_header = budget_table.rows[0].cells
    budget_header[0].text = 'Atribut'
    budget_header[1].text = 'Tip'
    budget_header[2].text = 'Descriere'
    
    for cell in budget_header:
        set_cell_background(cell, 'F7DC6F')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(51, 51, 51)
    
    budget_data = [
        ('Id', 'int', 'Cheie primară, auto-increment'),
        ('Category', 'string', 'Categoria pentru care se setează bugetul'),
        ('MonthlyLimit', 'double', 'Limita bugetului lunar în RON'),
        ('Month', 'int', 'Luna (1-12)'),
        ('Year', 'int', 'Anul bugetului')
    ]
    
    for idx, (attr, typ, desc) in enumerate(budget_data):
        row = budget_table.rows[idx + 1]
        row.cells[0].text = attr
        row.cells[1].text = typ
        row.cells[2].text = desc
        
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, 'FFFEF0')

def create_services(doc):
    """Secțiunea 4: Servicii de Bază"""
    add_page_break(doc)
    add_heading(doc, '4. SERVICII DE BAZĂ', level=1)
    
    add_heading(doc, '4.1. DatabaseService', level=2)
    add_normal_text(doc,
        'DatabaseService este serviciul central pentru toate operațiile cu baza de date SQLite. '
        'Oferă acces async la operațiile CRUD (Create, Read, Update, Delete) pentru toate entityurile.')
    
    add_heading(doc, 'Metode Principale:', level=3)
    
    db_methods = [
        ('InitAsync()', 'Inițializează conexiunea, creează tabele și face seed cu date implicite'),
        ('GetExpensesAsync()', 'Returnează toate cheltuielile sortate descrescător după dată'),
        ('GetExpensesByDateRangeAsync(start, end)', 'Returnează cheltuielile din interval'),
        ('SaveExpenseAsync(expense)', 'Salvează cheltuiala nouă sau actualizează cea existentă'),
        ('DeleteExpenseAsync(expense)', 'Șterge cheltuiala din baza de date'),
        ('GetCategoriesAsync()', 'Returnează toate categoriile disponibile'),
        ('SaveCategoryAsync(category)', 'Salvează categorie nouă sau actualizează cea existentă'),
        ('DeleteCategoryAsync(category)', 'Șterge categorie din baza de date'),
        ('GetBudgetsAsync()', 'Returnează toate bugetele'),
        ('GetBudgetsForMonthAsync(month, year)', 'Returnează bugetele pentru luna specificată'),
    ]
    
    for method_name, description in db_methods:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f'{method_name}: ')
        run.font.bold = True
        run.font.size = Pt(10)
        run = p.add_run(description)
        run.font.size = Pt(10)
    
    add_heading(doc, '4.2. CurrencyService', level=2)
    add_normal_text(doc,
        'CurrencyService gestionează conversiile de valute prin API extern (exchangerate-api.com). '
        'Oferă rate de schimb actualizate și conversii precise între valute.')
    
    add_heading(doc, 'Metode:', level=3)
    
    curr_methods = [
        ('GetExchangeRatesAsync()', 'Apelează API-ul pentru rate curente, cu fallback la rate implicite'),
        ('ConvertCurrencyAsync(amount, from, to)', 'Convertește suma de la o valută la alta')
    ]
    
    for method_name, description in curr_methods:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f'{method_name}: ')
        run.font.bold = True
        run.font.size = Pt(10)
        run = p.add_run(description)
        run.font.size = Pt(10)
    
    add_heading(doc, '4.3. Valute Suportate', level=3)
    
    currencies = [
        'RON (Leu Român) - Valuta implicită, bază de referință',
        'EUR (Euro) - Rata implicită: 0.20',
        'USD (Dolar American) - Rata implicită: 0.22',
        'GBP (Lira Sterlină) - Rata implicită: 0.17'
    ]
    
    for curr in currencies:
        p = doc.add_paragraph(curr)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)

def create_ui_pages(doc):
    """Secțiunea 5: Interface-uri Utilizator"""
    add_page_break(doc)
    add_heading(doc, '5. INTERFACE-URI UTILIZATOR', level=1)
    add_normal_text(doc,
        'Spending Tracker dispune de 9 pagini principale, fiecare cu o funcție specifică și ușor accesibilă din meniu lateral.')
    
    pages_info = [
        ('MainPage', 'Pagina Principală', 
         'Punct de intrare principal. Permite adăugarea rapidă a cheltuielilor cu toți parametrii necesari.',
         'MainPage - Adăugare Cheltuiala'),
        
        ('ExpenseListPage', 'Lista Cheltuieli',
         'Afișează toate cheltuielile cu posibilitate de filtrare pe categorii. Cheltuielile sunt sortate descrescător după dată.',
         'ExpenseListPage - Lista Cheltuieli Filtrate'),
        
        ('BudgetPage', 'Gestionare Bugete',
         'Permite setarea și monitorizare bugetelor lunare per categorie. Afișează ProgressBar cu procent de utilizare.',
         'BudgetPage - Monitorizare Bugete'),
        
        ('CategoriesPage', 'Gestionare Categorii',
         'CRUD complet pentru categorii. Permite adăugarea de categorii noi cu culoare personalizată și ștergerea celor existente.',
         'CategoriesPage - Gestionare Categorii'),
        
        ('StatisticsPage', 'Statistici',
         'Analiză vizuală a cheltuielilor cu progrese grafice. Afișează: total lunar, media zilnică, distribuție pe categorii.',
         'StatisticsPage - Analiză Statistică'),
        
        ('ReportsPage', 'Rapoarte',
         'Rapoarte detaliate pe perioade diferite. Conține: total, media zilnică, ziua cu cheltuieli maxime, top 5 categorii.',
         'ReportsPage - Rapoarte Periodice'),
        
        ('CurrencyConverterPage', 'Convertor Valute',
         'Conversie monede în timp real cu rate actualizate de la API. Suportă RON, EUR, USD, GBP.',
         'CurrencyConverterPage - Convertor Valute'),
        
        ('SettingsPage', 'Setări',
         'Configurări ale aplicației. Permite setarea bugetelor și afișează statistici generale.',
         'SettingsPage - Setări Aplicație'),
        
        ('AboutPage', 'Despre',
         'Informații despre aplicație, versiune, ghid utilizare, funcționalități principale și informații tehnice.',
         'AboutPage - Despre Aplicație')
    ]
    
    for page_file, page_title, description, image_name in pages_info:
        add_heading(doc, f'5.{pages_info.index((page_file, page_title, description, image_name)) + 1}. {page_title}', level=2)
        add_normal_text(doc, description)
        add_image_placeholder(doc, image_name, f'Screenshot-ul paginii {page_title}')
        doc.add_paragraph()

def create_functionalities(doc):
    """Secțiunea 6: Funcționalități Principale"""
    add_page_break(doc)
    add_heading(doc, '6. FUNCȚIONALITĂȚI PRINCIPALE', level=1)
    
    add_heading(doc, '6.1. Adăugare și Editare Cheltuieli', level=2)
    add_normal_text(doc,
        'Procesul de adăugare este simplu și rapid. Utilizatorul completează formularul cu descrierea, suma, categoria, '
        'data și valuta. Aplicația validează datele și le stochează în baza de date SQLite.')
    
    add_heading(doc, 'Validări aplicate:', level=3)
    validations = [
        'Descriere: Nu poate fi goală, lungime minimă 3 caractere',
        'Suma: Trebuie număr pozitiv, format 0.00',
        'Categoria: Trebuie selectată din listă',
        'Dată: Implicit azi, poate fi editată',
        'Valută: Implicit RON, opțiuni: RON, EUR, USD, GBP'
    ]
    
    for validation in validations:
        p = doc.add_paragraph(validation)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_image_placeholder(doc, 'AddExpense_Workflow', 'Fluxul de adăugare a unei cheltuieli')
    
    add_heading(doc, '6.2. Filtrare și Căutare', level=2)
    add_normal_text(doc,
        'ExpenseListPage permite filtrarea cheltuielilor pe categorii. La selectie categoria din Picker, '
        'lista se actualizează automat afișând doar cheltuielile din categoria respectivă.')
    
    add_heading(doc, '6.3. Gestionare Bugete Lunare', level=2)
    add_normal_text(doc,
        'Utilizatorul poate seta bugete lunare pentru fiecare categorie. Aplicația calculează automat '
        'cheltuielile lunii și afișează procentul de utilizare cu indicator culoare: Verde < 80%, Portocaliu 80-100%, Roșu > 100%.')
    
    add_image_placeholder(doc, 'Budget_Monitoring', 'Monitorizare bugete cu indicatori de culoare')
    
    add_heading(doc, '6.4. Analiză Statistică', level=2)
    add_normal_text(doc,
        'StatisticsPage calculează și afișează metrici importante: total lunar, media zilnică, distribuție pe categorii. '
        'Utilizatorul poate selecta interval de date pentru analiză personalizată.')
    
    add_heading(doc, '6.5. Rapoarte Periodice', level=2)
    add_normal_text(doc,
        'ReportsPage oferă rapoarte detaliate pe 5 perioade predefinite: luna curentă, luna trecută, '
        'ultimele 3 luni, ultimele 6 luni, ultimul an.')
    
    add_heading(doc, '6.6. Conversia Valutelor', level=2)
    add_normal_text(doc,
        'Convertor profesional cu rate actualizate zilnic din API extern. Dacă API nu e disponibil, '
        'se folosesc rate implicite pentru continuitate.')
    
    add_image_placeholder(doc, 'Currency_Converter', 'Interfața convertorului de valute')

def create_flows(doc):
    """Secțiunea 7: Fluxuri de Date"""
    add_page_break(doc)
    add_heading(doc, '7. FLUXURI DE DATE ȘI PROCESE', level=1)
    
    add_heading(doc, '7.1. Flux de Adăugare Cheltuială', level=2)
    flow1 = """
1. Utilizator introduce date în MainPage
2. OnSaveClicked() event handler
3. Validare date (descriere, suma, categorie)
4. Creare obiect Expense cu datele introduse
5. _databaseService.SaveExpenseAsync(expense)
6. DatabaseService.InitAsync() - inițializare DB
7. _database.InsertAsync(expense) - inserare în SQLite
8. Returare ID nou generat
9. Afișare mesaj de succes
10. Resetare form
    """
    add_normal_text(doc, flow1)
    
    add_heading(doc, '7.2. Flux de Analiză Statistică', level=2)
    flow2 = """
1. StatisticsPage.OnAppearing() - pagina e accesată
2. LoadStatistics() - încarcă date din DB
3. GetExpensesByDateRangeAsync(start, end)
4. Query: SELECT * FROM Expense WHERE Date BETWEEN start/end
5. Calcul metrici: Total, Count, Average, GroupBy Category
6. Creare ViewModels pentru afișare
7. Set ItemsSource pe CollectionView
8. UI se actualizează - ProgressBars se umple
    """
    add_normal_text(doc, flow2)
    
    add_heading(doc, '7.3. Flux de Gestionare Bugete', level=2)
    flow3 = """
1. Utilizator accesează BudgetPage
2. OnAppearing() apelează LoadBudgets()
3. GetBudgetsForMonthAsync(month, year) - budgete luna curentă
4. Pentru fiecare buget se calculează cheltuielile din categoria respectivă
5. Calcul progres: Progress = Cheltuieli / MonthlyLimit (maxim 1.0)
6. Determinare culoare: Verde < 0.8, Portocaliu 0.8-1.0, Roșu > 1.0
7. Creare BudgetViewModel cu datele calculate
8. UI afișează ProgressBars colorate corespunzător
    """
    add_normal_text(doc, flow3)
    
    add_image_placeholder(doc, 'Data_Flow_Diagram', 'Diagrama fluxului de date al sistemului')

def create_security(doc):
    """Secțiunea 8: Securitate"""
    add_page_break(doc)
    add_heading(doc, '8. CARACTERISTICI DE SECURITATE', level=1)
    
    add_heading(doc, '8.1. Validare Date', level=2)
    add_normal_text(doc,
        'Toate input-urile sunt validate atât la nivel de UI cât și la nivel de business logic. '
        'Se verifică tipuri de date, valori permise și lungimi minime/maxime.')
    
    add_heading(doc, '8.2. Gestionare Erori', level=2)
    add_normal_text(doc,
        'Toate operațiile async sunt înconjurate cu try-catch blocks. Se loghează erori pentru debugging '
        'și utilizatorul e informat prin DisplayAlert.')
    
    add_heading(doc, '8.3. Confirmări Critice', level=2)
    add_normal_text(doc,
        'Operații critice (ștergere date) necesită confirmare explicită din utilizator. '
        'Mesajele de confirmare sunt clare și avertizează despre consecințe.')
    
    critical_ops = [
        'Ștergerea tuturor cheltuielilor: "Ești sigur că vrei să ștergi toate cheltuielile? Această acțiune nu poate fi anulată!"',
        'Ștergerea categorie: "Sigur vrei să ștergi această categorie?"'
    ]
    
    for op in critical_ops:
        p = doc.add_paragraph(op)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '8.4. Bază de Date Locală', level=2)
    add_normal_text(doc,
        'SQLite stochează datele pe dispozitiv în format criptat. Nu sunt trimise date la servere externe. '
        'Utilizatorul are control total asupra datelor sale.')

def create_user_guide(doc):
    """Secțiunea 9: Ghid Utilizator"""
    add_page_break(doc)
    add_heading(doc, '9. GHID UTILIZATOR', level=1)
    
    add_heading(doc, '9.1. Primii Pași', level=2)
    
    steps = [
        'Descarcă și instalează Spending Tracker din Google Play Store sau App Store',
        'La prima lansare, aplicația creează baza de date și categoriile implicite',
        'Sunt preîncărcate 7 categorii: Alimente, Transport, Divertisment, Utilități, Sănătate, Educație, Altele'
    ]
    
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{idx}. {step}')
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '9.2. Adăugare Cheltuiala - Tutorial', level=2)
    
    tutorial = [
        '1. Accesează "Acasă" din meniu lateral',
        '2. Introdu descrierea (ex: "Cumpărături Lidl")',
        '3. Introdu suma (ex: 85.50)',
        '4. Selectează categoria (ex: "Alimente")',
        '5. Data e implicit astazi, poți schimba dacă vrei',
        '6. Selectează valuta (implicit RON)',
        '7. Apasă "Salvează"',
        '8. Se afișează confirmare "Cheltuiala salvată cu succes"'
    ]
    
    for step in tutorial:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '9.3. Consulare Statistici - Tutorial', level=2)
    
    tutorial2 = [
        '1. Accesează "Statistici" din meniu',
        '2. Selectează dată inițială și finală cu DatePicker',
        '3. Observă: Total, Media zilnică, Distribuție categorii',
        '4. ProgressBar-urile arată procent din total pe fiecare categorie',
        '5. Colori indicator: Verde = bine, Portocaliu = atent, Roșu = depășit'
    ]
    
    for step in tutorial2:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '9.4. Setare Bugete - Tutorial', level=2)
    
    tutorial3 = [
        '1. Accesează "Bugete" din meniu',
        '2. Selectează categoria pentru care vrei buget',
        '3. Introdu suma bugetului lunar (ex: 500 RON)',
        '4. Apasă "Salvează Buget"',
        '5. Vizualizează bugetele active cu procent utilizat',
        '6. Apasă "Rapoarte" pentru a vedea cum stai pe fiecare categorie'
    ]
    
    for step in tutorial3:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)

def create_conclusions(doc):
    """Secțiunea 10: Concluzii"""
    add_page_break(doc)
    add_heading(doc, '10. CONCLUZII ȘI PERSPECTIVĂ VIITOARE', level=1)
    
    add_normal_text(doc,
        'Spending Tracker oferă o soluție completă, ușor de utilizat, pentru gestionarea cheltuielilor zilnice. '
        'Cu 9 pagini funcționale și bază de date locală sigură, aplicația permite utilizatorilor să înțeleagă mai '
        'bine obiceiurile lor de cheltuire și să stabilească bugete realiste.')
    
    add_heading(doc, '10.1. Puncte Forte Actuale', level=2)
    
    strengths = [
        '✓ Interfață intuitivă și ușor de utilizat (designul modern și responsive)',
        '✓ Suport multiplatformă (Android, iOS, Windows, macOS)',
        '✓ Bază de date locală sigură (SQLite cu date criptate pe dispozitiv)',
        '✓ Statistici și rapoarte detaliate cu vizualizare progrese',
        '✓ Gestionare bugete eficientă cu alerte vizuale colorate',
        '✓ Convertor de valute cu rate în timp real',
        '✓ Validare riguroasă a datelor și gestionare erori',
        '✓ Cod modular și bine structurat (ușor de menținut și extins)',
        '✓ Performance excelent (bază de date locală)',
        '✓ Fără reclame sau colectare date personale'
    ]
    
    for strength in strengths:
        p = doc.add_paragraph(strength)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '10.2. Funcționalități Viitoare Planificate', level=2)
    
    future = [
        '→ Export rapoarte (PDF, Excel, CSV)',
        '→ Sincronizare cloud (OneDrive, Google Drive) cu backup automat',
        '→ Notificări push de avertizare buget la depășire',
        '→ Grafice avansate cu Microcharts și animații',
        '→ Backup și restore automat periodic',
        '→ Sarcini recurente pentru cheltuieli fixe',
        '→ Colaborare între utilizatori în familii',
        '→ Predictare cheltuieli pe bază de IA/ML',
        '→ Integrare cu bănci pentru importare automată tranzacții',
        '→ Notificări SMS/Email pentru alerte importante'
    ]
    
    for feature in future:
        p = doc.add_paragraph(feature)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '10.3. Recomandări de Optimizare', level=2)
    
    recommendations = [
        'Cache-are date frecvent accesate pentru performanță mai bună',
        'Paginare pentru liste lungi (> 1000 elemente)',
        'Search full-text pentru cheltuieli cu filtrare avansată',
        'Sincronizare offline-first cu reconciliere cloud',
        'Testare A/B pentru UI improvements',
        'Analytics pentru înțelegere comportament utilizatori'
    ]
    
    for rec in recommendations:
        p = doc.add_paragraph(rec)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    add_heading(doc, '10.4. Concluzie Finală', level=2)
    add_normal_text(doc,
        'Spending Tracker este o aplicație de producție, gata pentru utilizare în masă, care oferă valoare reală '
        'utilizatorilor în gestionarea finanțelor personale. Codul este bine structurat, extensibil și ușor de menținut, '
        'permițând adăugări ușoare de noi funcționalități în viitor. Aplicația răspunde unei nevoi reale pe piață și are '
        'potențial semnificativ de creștere și monetizare.')
    
    add_normal_text(doc,
        'Cu o bază de utilizatori dedicată și feedback pozitiv, Spending Tracker ar putea deveni o soluție lider în '
        'gestionarea cheltuielilor personale pe piața mobilă din România și nu numai.')
    
    # Final info
    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run('© 2025 Spending Tracker. Toate drepturile rezervate.')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    final2 = doc.add_paragraph()
    final2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final2.add_run(f'Documentație generată: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

def create_word_documentation():
    """Funcția principală de creație a documentului"""
    print("=" * 70)
    print("SPENDING TRACKER - GENERATOR DOCUMENTAȚIE WORD")
    print("=" * 70)
    print("\n📄 Se generează documentul Word...")
    print("Așteptați, procesul este în curs...\n")
    
    # Creare document
    doc = Document()
    
    # Setare marje
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Pagina de titlu
    print("✓ Creez pagina de titlu...")
    create_title_page(doc)
    add_page_break(doc)
    
    # Cuprins
    print("✓ Creez cuprinsul...")
    create_toc(doc)
    add_page_break(doc)
    
    # Secțiuni principale
    print("✓ Creez secțiunea 1: Introducere...")
    create_introduction(doc)
    
    print("✓ Creez secțiunea 2: Descriere Tehnică...")
    create_technical_description(doc)
    
    print("✓ Creez secțiunea 3: Modelul de Date...")
    create_data_model(doc)
    
    print("✓ Creez secțiunea 4: Servicii de Bază...")
    create_services(doc)
    
    print("✓ Creez secțiunea 5: Interface-uri Utilizator...")
    create_ui_pages(doc)
    
    print("✓ Creez secțiunea 6: Funcționalități Principale...")
    create_functionalities(doc)
    
    print("✓ Creez secțiunea 7: Fluxuri de Date...")
    create_flows(doc)
    
    print("✓ Creez secțiunea 8: Securitate...")
    create_security(doc)
    
    print("✓ Creez secțiunea 9: Ghid Utilizator...")
    create_user_guide(doc)
    
    print("✓ Creez secțiunea 10: Concluzii...")
    create_conclusions(doc)
    
    # Salvare document
    filename = "SPENDING_TRACKER_DOCUMENTATION.docx"
    doc.save(filename)
    
    print("\n" + "=" * 70)
    print("✅ SUCCES! Documentația Word a fost generată cu succes!")
    print("=" * 70)
    print(f"\n📄 Fișierul: {filename}")
    print(f"📊 Dimensiune: {os.path.getsize(filename) / 1024 / 1024:.2f} MB")
    print(f"📅 Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    print("\n📝 Documentul conține:")
    print("  ✓ Pagina de titlu profesională")
    print("  ✓ Cuprins cu 10 secțiuni")
    print("  ✓ Descriere tehnică completă")
    print("  ✓ Tabele colorate cu date")
    print("  ✓ 9 pagini dedicate interface-urilor cu PLACEHOLDER-URI PENTRU IMAGINI")
    print("  ✓ Fluxuri de date și procese")
    print("  ✓ Caracteristici de securitate")
    print("  ✓ Ghid utilizator detaliat")
    print("  ✓ Concluzii și perspectivă viitoare")
    
    print("\n🎨 Placeholder-uri pentru imagini (marcate în RED):")
    print("  [INSERT IMAGE: MainPage - Adăugare Cheltuiala]")
    print("  [INSERT IMAGE: ExpenseListPage - Lista Cheltuieli Filtrate]")
    print("  [INSERT IMAGE: BudgetPage - Monitorizare Bugete]")
    print("  [INSERT IMAGE: CategoriesPage - Gestionare Categorii]")
    print("  [INSERT IMAGE: StatisticsPage - Analiză Statistică]")
    print("  [INSERT IMAGE: ReportsPage - Rapoarte Periodice]")
    print("  [INSERT IMAGE: CurrencyConverterPage - Convertor Valute]")
    print("  [INSERT IMAGE: SettingsPage - Setări Aplicație]")
    print("  [INSERT IMAGE: AboutPage - Despre Aplicație]")
    print("  [INSERT IMAGE: AddExpense_Workflow]")
    print("  [INSERT IMAGE: Budget_Monitoring]")
    print("  [INSERT IMAGE: Currency_Converter]")
    print("  [INSERT IMAGE: Data_Flow_Diagram]")
    
    print("\n💡 Instrucțiuni de completare:")
    print("  1. Deschide fișierul Word")
    print("  2. Caută [INSERT IMAGE: ...]")
    print("  3. Selectează textul placeholder")
    print("  4. Șterge-l și inserează imaginea reală")
    print("  5. Redimensionează imaginea dacă necesare")
    print("  6. Repetă pentru toate placeholder-urile")
    
    print("\n✅ Documentul este gata pentru completare cu imagini!")
    print("=" * 70 + "\n")

if __name__ == "__main__":
    try:
        import os
        create_word_documentation()
    except ImportError as e:
        print(f"\n❌ EROARE: Lipsă dependență Python")
        print(f"Detaliu: {e}")
        print("\nInstalează dependențele cu comanda:")
        print("  pip install python-docx Pillow")
    except Exception as e:
        print(f"\n❌ EROARE: {e}")
        import traceback
        traceback.print_exc()

